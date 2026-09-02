import json
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "source_data" / "roguelike_topic_table_full.json"
OUT_JSON = ROOT / "黑流树海模拟器规则库.json"
OUT_DOCX = ROOT / "黑流树海节点事件与收益模拟器规范.docx"
ACCESSED = "2026-08-31"


def clean(text):
    if text is None:
        return ""
    text = str(text).replace("\\n", "；")
    text = re.sub(r"<[^>]+>", "", text)
    return text.replace("</>", "").strip()


def event(name, node, availability, options, repeat=False, confidence="B", notes=""):
    return {
        "name": name,
        "node_type": node,
        "availability": availability,
        "repeatable": repeat,
        "confidence": confidence,
        "options": options,
        "notes": notes,
    }


def option(label, requirement="始终", cost="无", result="", transition="结束事件", confidence="B"):
    return {
        "label": label,
        "requirement": requirement,
        "cost": cost,
        "result": result,
        "transition": transition,
        "confidence": confidence,
    }


EVENTS = [
    event("桑尼的邀请", "INCIDENT", "已确认可在第I层出现", [
        option("帮新郎唱首歌", result="希望+3"),
        option("收下谢礼", result="源石锭+10"),
    ]),
    event("色味不同源", "INCIDENT", "已确认可在第I层出现", [
        option("吃下苹果", result="目标生命上限+3"),
        option("吃下鱼", result="护盾+5"),
        option("离开"),
    ]),
    event("货从口出", "INCIDENT", "楼层池未公开", [
        option("购买", cost="源石锭-4", result="随机获得1个零件；类型可为加工品/概念体/自然物", transition="可再次购买或离开"),
        option("离开"),
    ], repeat=True, notes="零件类型候选已知，类型及物品权重未知。"),
    event("沉重的契约", "INCIDENT", "楼层池未公开", [
        option("签下契约", result="行动力+2"),
        option("以血签约", cost="当前目标生命值减半", result="行动力+5"),
        option("离开"),
    ]),
    event("敲动杠杆", "INCIDENT", "已确认可在第I层出现", [
        option("轻轻敲动", cost="源石锭-4", result="随机收藏品+1", transition="可继续或离开"),
        option("用力敲动", cost="源石锭-10", result="随机收藏品+3", transition="可继续或离开"),
        option("离开"),
    ], repeat=True, notes="收藏品抽取权重未公开。"),
    event("血衣之下", "INCIDENT", "楼层池未公开", [
        option("选择『他缚』", result="获得对应事件收藏品『他缚』"),
        option("选择『剑锤』", result="获得对应事件收藏品『剑锤』"),
        option("选择『一串钱伥』", result="获得对应事件收藏品『一串钱伥』"),
        option("转身离开", result="目标生命上限+2"),
    ], notes="三个事件收藏品的客户端ID由名称解析表给出。"),
    event("擒与缚", "INCIDENT", "楼层池未公开", [
        option("选择『翱翼』", result="获得『翱翼』"),
        option("选择『虬蜕』", result="获得『虬蜕』"),
        option("离开"),
    ]),
    event("沉寂之屋", "INCIDENT", "已确认可在第I层出现", [
        option("以血换取", requirement="目标生命值足够支付", cost="目标生命值-2", result="获得零件『笼控器』"),
        option("购买", requirement="源石锭≥8", cost="源石锭-8", result="获得零件『笼控器』"),
        option("离开"),
    ]),
    event("黑诞", "INCIDENT", "楼层池未公开；与『沉寂之屋』支线关联", [
        option("出示笼控器", requirement="持有『笼控器』", result="获得『猎印』"),
        option("迎战", result="随机进入『闹乐』『纵怒』『灭身』之一；胜利结算特殊战斗奖励", transition="战斗结算"),
        option("逃离", cost="目标生命值-3（最低保留1）", result="无额外奖励"),
    ], notes="三个战斗关卡的抽取权重未知。"),
    event("呼吸的红苔", "INCIDENT", "需先经历『沉寂之屋』；楼层池未公开", [
        option("使用笼控器", requirement="持有『笼控器』", result="源石锭+30；下一层出现希望地块"),
        option("播种", requirement="持有可培育的『种子』", cost="消耗种子", result="执行种子培育结果；下一层出现希望地块"),
        option("离开"),
    ]),
    event("被歌颂的影子", "INCIDENT", "楼层池未公开", [
        option("第一圈", cost="目标生命值-1", result="随机：源石锭+4，或回复1目标生命", transition="进入下一圈或离开"),
        option("第二圈", cost="目标生命值-1", result="随机：希望+2，或回复1目标生命", transition="进入下一圈或离开"),
        option("第三圈", cost="目标生命值-2", result="随机自然物+1，或回复2目标生命", transition="进入下一圈或离开"),
        option("第四圈", cost="目标生命值-2", result="随机：加工品+1、源石锭+12、或回复2目标生命", transition="进入下一圈或离开"),
        option("第五圈", cost="目标生命值-2", result="随机收藏品+1，或回复5目标生命", transition="进入下一圈或离开"),
        option("第六圈", cost="目标生命值-3", result="随机：希望+8、编队上限+1、或回复3目标生命", transition="进入下一圈或离开"),
        option("第七圈", cost="目标生命值-3", result="随机收藏品+3", transition="进入下一圈或离开"),
        option("第八圈", cost="目标生命值-3", result="未持有时获得『犬植浆』；已持有时行动力+1"),
        option("离开", result="保留此前已取得的结果"),
    ], notes="每圈分支概率未公开；应由外部权重表注入。"),
    event("愈创之心", "INCIDENT", "楼层池未公开；结局支线事件", [
        option("以行动力交换", requirement="行动力≥3", cost="行动力-3", result="获得『追忆』"),
        option("献上加工品", requirement="加工品≥2", cost="随机消耗2个加工品", result="获得『追忆』"),
        option("使用源私钥", requirement="持有『源私钥』", cost="消耗『源私钥』", result="获得『追忆』"),
        option("启动笼控器", requirement="持有『笼控器』", result="行动力+5"),
        option("休整", result="行动力+1"),
    ]),
    event("思乡心切", "INCIDENT", "已确认可在第I层出现", [
        option("接受挑战", result="进入特殊战斗『陌生旅伴』", transition="战斗结算"),
        option("离开"),
    ]),
    event("划算买卖", "INCIDENT", "楼层池未公开", [
        option("处理问题", result="进入特殊战斗『安保措施』", transition="战斗结算"),
        option("离开"),
    ]),
    event("鸭托邦", "INCIDENT", "楼层池未公开", [
        option("嘲弄", result="进入普通难度『开业剪彩』", transition="战斗结算"),
        option("挑衅", result="进入高难/紧急难度『开业剪彩』", transition="战斗结算"),
        option("离开"),
    ]),
    event("传奇团伙", "INCIDENT", "楼层池未公开", [
        option("入座", result="进入普通难度『合伙人会议』", transition="战斗结算"),
        option("戳破骗局", requirement="源石锭>50", result="进入紧急难度『合伙人会议』", transition="战斗结算"),
        option("交钱脱身", cost="失去当前源石锭的一半", result="离开事件"),
    ]),
    event("湖中仙女", "INCIDENT", "楼层池未公开", [
        option("献上供物（第1—3次）", requirement="每次源石锭≥1", cost="每次源石锭-1", result="推进供奉；前三次均可选择离开", transition="继续供奉或离开"),
        option("揭穿倒影", result="进入普通『湖中魇』", transition="战斗结算"),
        option("完成三次供奉", result="随机：进入紧急『湖中魇』，或获得珍贵收藏品"),
        option("离开"),
    ], notes="第三次供奉的分支概率与珍贵收藏品候选池未公开。"),
    event("洞中宝", "INCIDENT", "楼层池未公开", [
        option("挖掘", result="随机：源石锭+5、收藏品+1、无事发生并可继续、或遭咬伤并使下一层普通作战变为高难", transition="依结果继续或结束"),
        option("离开"),
    ], repeat=True, notes="公开页面未给出完整递进轮次和权重；客户端选项原文已附在JSON原始目录。"),
    event("临时中介所", "INCIDENT", "楼层池未公开", [
        option("聘用王牌", result="获得随机六星临时招募干员"),
        option("聘用熟手", result="希望+2；从先锋/近卫、重装/辅助、医疗/狙击、术师/特种之一的组合中招募已晋升的五星及以下干员"),
        option("离开"),
    ], notes="职业组合及干员抽取权重未公开。"),
    event("和平守卫者", "INCIDENT", "楼层池未公开", [
        option("接下武器", result="获得『厄运火杆』；若区域内有未完成的最近居民据点则立即前往，否则标记最近据点"),
        option("离开"),
    ]),
    event("独活", "INCIDENT", "社区资料标注第II层；需客户端/实测复核", [
        option("接受同行", result="获得『同行者』；立即前往最近的紧急作战协助，或标记目标"),
        option("离开"),
    ], confidence="C"),
    event("线人", "INCIDENT", "结局路线资料标注第II—IV层；需实测复核", [
        option("接头", result="获得『α』与1个稀有加工品"),
        option("离开", transition="未进入核心选项时事件可保留/再次访问"),
    ], repeat=True, confidence="C"),
    event("泪之聚落", "INCIDENT", "结局路线资料标注第II—IV层；需实测复核", [
        option("倾尽所有", cost="失去全部源石锭", result="获得『击坠神明』"),
        option("离开", transition="未进入核心选项时事件可保留/再次访问"),
    ], repeat=True, confidence="C"),
    event("安眠一隅", "INCIDENT_HIDDEN", "不会自然生成；使用『小八界』时有极低概率传送", [
        option("接受馈赠", result="获得『源私钥』（其效果包含源石锭+99）；立即进入下一层并触发乌托邦『美丽新大地』"),
    ], confidence="B", notes="传送概率未公开。"),
    event("未涉足之树", "EXPEDITION", "除第I、V、VI层外", [
        option("尝试探索树的内部", requirement="选择1名已招募干员", cost="该干员离队至下一层", result="下一层归队；常规路线按客户端远征结算；持『生命游戏·喙』时额外获得随机加工品"),
        option("探索源头", requirement="选择1名已招募干员", cost="该干员离队至下一层", result="归队时不晋升；希望+2；获得『怦然』信号"),
        option("在树下休息", result="希望+2"),
    ], notes="常规路线的晋升细节应以实际远征结算配置为准。"),
    event("回滚文明", "SACRIFICE", "除第I、VI层外", [
        option("置换收藏品", requirement="持有可献祭收藏品", cost="消耗所选收藏品", result="按稀有度规则获得随机收藏品；持『生命游戏·手』时可追加1次同类型交换"),
        option("置换零件", requirement="持有零件且声带能力开放", cost="消耗所选零件", result="按零件稀有度规则获得随机零件；持『生命游戏·手』时可追加1次同类型交换"),
        option("燃尽", requirement="持有『怦然』且自然物≥2", cost="随机消耗2个自然物", result="获得『焚毁文明』"),
        option("离开"),
    ]),
    event("无人商店", "WISH", "得偿所愿的特定场景", [
        option("打开大木桶", result="从该槽位展示的收藏品中选择/获得1个"),
        option("打开两个小木桶", result="从该槽位展示的收藏品中选择/获得1个"),
        option("打开巨型带轮容器", requirement="持有『四叶草化石』", result="获得该隐藏槽位收藏品"),
        option("撬开补货", requirement="源石锭≥4且未刷新", cost="源石锭-4", result="刷新一次，并提高展示收藏品档位"),
        option("离开"),
    ], notes="具体槽位候选池和权重未在公开客户端表中完整暴露。"),
    event("溯源", "PORTAL", "除第I、II、VI层外；六种目的地变体", [
        option("投入加工品", requirement="至少持有界面列出的1个加工品", cost="消耗所选加工品", result="进入对应黑色水池/特殊层；加工品在进入和离开时不额外损耗"),
        option("离开"),
    ], notes="界面最多列出3个可用加工品；六变体目的地映射仍需实测/服务器权重。"),
    event("原始娱乐", "DUEL", "狭路相逢；不在第I、II、VI层出现", [
        option("应战", result="公平作战；失败不扣目标生命/护盾；胜利获得低/中稀有零件+2张招募券"),
        option("离开"),
    ]),
    event("掠夺成性", "DUEL", "狭路相逢；不在第I、II、VI层出现", [
        option("应战", result="公平作战；失败不扣目标生命/护盾；胜利获得中/高稀有零件或『浪花/雾滚草/小八界』之一+2张招募券"),
        option("离开"),
    ]),
    event("好奇心与死", "STORY", "结局链节点", [
        option("追踪目标", requirement="源石锭≥50，且未同时满足两处沙坑/未访问『窥视箱中』等互斥条件", cost="源石锭-50", result="标记或定位目标节点"),
        option("取走收藏品", result="获得场景收藏品"),
    ], notes="互斥条件需由结局状态机维护。"),
    event("窥视箱中", "STORY", "结局链节点", [
        option("保护箱中之物", result="节点转化为首领战『混沌源阶理论』；转化后不可离开", transition="战斗结算"),
        option("离开", transition="保留事件状态"),
    ]),
    event("调谐仪式", "STORY", "结局链节点", [
        option("提交击坠神明", requirement="持有『击坠神明』或用『源私钥』替代", cost="消耗提交物", result="首个『灾厄之口』不再摧毁『危朽』"),
        option("提交焚毁文明", requirement="持有『焚毁文明』或用『源私钥』替代", cost="消耗提交物", result="不再锁定部署区"),
        option("提交湮没光明", requirement="持有『湮没光明』或用『源私钥』替代", cost="消耗提交物", result="降低最终战相关伤害"),
        option("不提交", result="不获得对应削弱"),
    ]),
    event("险路尽头", "FINAL", "区域出口；可重复访问，信息默认揭示", [
        option("进入下一区域", result="随机加工品+1；剩余行动力等量转化为希望；可领取寄存招募券；进入下一层"),
    ], repeat=True, notes="保密度≥3时会少揭示1项信息。"),
    event("三重身·年幼", "EVACUATE", "险路小径出口变体", [option("撤离", result="获得稀有加工品；保留剩余行动力；不可领取寄存招募券；进入下一层")], repeat=True),
    event("三重身·壮年", "EVACUATE", "险路小径出口变体", [option("撤离", result="获得稀有加工品；保留剩余行动力；不可领取寄存招募券；进入下一层")], repeat=True),
    event("三重身·老去", "EVACUATE", "险路小径出口变体", [option("撤离", result="获得稀有加工品；保留剩余行动力；不可领取寄存招募券；进入下一层")], repeat=True),
    event("金色凝滞", "REST", "安全的角落场景；第I层不出现", [
        option("从3项展示中选1项", result="候选池6项：目标生命上限+3、编队上限+1、希望+3、高级补给券+1、行动力+2、零件箱容量+1"),
    ], notes="每次从6项随机展示3项；展示权重未公开。"),
]


BATTLE_REWARDS = [
    ["I", "普通", "10—13", "1"], ["I", "紧急", "12—15", "2"],
    ["II", "普通", "12—15", "2"], ["II", "紧急", "18—23", "2"],
    ["III", "普通", "13—17", "2"], ["III", "紧急", "25—33", "3"], ["III", "险路恶敌", "32—42", "5"], ["III", "追猎", "32—42", "5"],
    ["IV", "普通", "15—19", "2"], ["IV", "紧急", "30—39", "3"],
    ["V", "普通", "20—26", "2"], ["V", "紧急", "36—47", "5"], ["V", "险路恶敌", "50—66", "8"], ["V", "追猎", "50—66", "8"],
    ["VI", "普通", "20—26", "5"], ["VI", "紧急", "36—47", "5"], ["VI", "险路恶敌", "70—93", "8"], ["VI", "追猎", "70—93", "8"],
]

SPECIAL_BATTLES = [
    ["思乡心切 / 陌生旅伴", "10—13", "2—3"], ["划算买卖 / 安保措施", "12—15", "2"],
    ["鸭托邦 / 开业剪彩", "13—17", "2"], ["开业剪彩（紧急）", "24—31", "3"],
    ["传奇团伙 / 合伙人会议", "25—33", "3"], ["合伙人会议（紧急）", "36—47", "3"],
    ["原始娱乐 / 搏杀", "13—17", "3"], ["掠夺成性 / 共斗", "18—23", "3"],
    ["行商挑衅 / 死箱", "10—13", "4"], ["行商挑衅 / 活木", "15—19", "5"],
    ["无效验尸", "36—47", "未确认"], ["湖中仙女 / 湖中魇", "25—33", "4"],
    ["湖中魇（紧急）", "36—47", "8"], ["流窜强买强卖", "15—19", "3"],
    ["进退趋同", "18—23", "5"], ["居民据点 / 枯枝", "25—33", "5"],
    ["居民据点 / 败叶", "30—39", "5"], ["黑诞三战", "30—39", "3"],
]


NODE_RULES = {
    "BATTLE_NORMAL": "随机普通作战；胜利结算当前层级经验与源石锭，并继续处理藏品/零件等战后额外掉落。",
    "BATTLE_ELITE": "随机紧急作战；奖励高于普通作战。",
    "BATTLE_BOSS": "险路恶敌；部分区域为强制出口。胜利后按表结算，离开区域时剩余行动力转希望。",
    "BATTLE_SHOP": "诡意行商：购买招募券、战术道具、收藏品、零件等。『生命游戏·脊椎』下可刷新4次，首次4锭、每次再+4。可在满足条件后挑衅行商。",
    "REST": "从随机展示的3项恢复/成长效果中选择1项；第I层不出现。",
    "INCIDENT": "从本规范事件池抽取并显示满足前置条件的选项；事件权重若未提供不得擅自均匀化。",
    "WISH": "通常展示2件收藏品免费二选一；可支付4锭刷新1次。特定场景为『无人商店』。",
    "SACRIFICE": "消耗收藏品进行同稀有度规则置换；声带能力开放时可换零件；『生命游戏·手』允许追加1次同类型交换。",
    "EXPEDITION": "派遣一名干员至下一层；未派遣则希望+2。『生命游戏·喙』使归队时额外获得随机加工品。",
    "PORTAL": "消耗一个界面允许的加工品进入特殊层/区域；第I、II、VI层不出现。",
    "DUEL": "公平作战，失败不扣目标生命或护盾；按左/中/右奖励档结算。第I、II、VI层不出现。",
    "STORY": "命运所指：按结局状态机触发；必须维护事件旗标和互斥条件。",
    "STORY_HIDDEN": "隐藏命运节点；生成条件依结局状态。",
    "SCRAP_SHOP": "秘境行商：买卖零件、培育种子；可重复访问。",
    "DOOR": "曲折密道成对传送；进入点消耗行动力，未完成时可越过；第I、II、VI层不出现。",
    "FINAL": "险路尽头：获得加工品，剩余行动力转希望，可领取寄存招募券，进入下一区域。",
    "EVACUATE": "险路小径：获得稀有加工品，保留行动力，不可领取寄存招募券，进入下一区域。",
    "EMPLOY": "应急助力：8名候选（通常5—6名普通费用4、2—3名本主题临时干员免费），最多招3名且均为已晋升；刷新费4起、每次+4、最多4次。",
    "LIGHT": "羽瞰点：初始揭示自身与周围4格；访问后扩为12格并行动力+1。保密度≥3且有『生命游戏』前额叶时扩为24格。",
    "BATTLE_SAVAGE": "『居民』据点：保密度≥4后出现；会生成游走居民。击破据点时清除区域内全部游走居民。第I、VI层不出现。",
    "EMPTY": "林间空地：不触发事件和收益。",
}


PART_ACQUISITION = {
    "报废轮子": "通用零件来源；险路恶敌/险路尽头可掉落", "报废假肢": "通用零件来源；险路恶敌/险路尽头可掉落",
    "标准引擎": "通用零件来源；险路恶敌/险路尽头可掉落", "气垫底座": "通用零件来源；险路恶敌/险路尽头可掉落",
    "试作外骨骼": "通用零件来源；险路恶敌/险路尽头可掉落", "重弹簧": "种子培育",
    "小八界": "种子培育；使用时极低概率传送至『安眠一隅』", "一次性喷气背包": "种子培育；险路小径",
    "老妈妈的融雪": "种子培育；险路小径", "坎诺特的触须": "种子培育；险路小径",
    "结构性原理": "不在行商出售；种子培育；险路小径", "“简易遥控器”": "不在行商出售；险路小径",
    "雾滚草": "种子培育", "回声玉米": "种子培育", "浪花": "种子培育", "霜晶树": "种子培育", "多生苔藓": "种子培育",
    "枯苔藓球": "不在行商出售", "板藤": "不在行商出售；藏果地", "恋家果": "不在行商出售；藏果地",
    "光彩松露": "不在行商出售；藏果地；『光荣远征』让赤金条进入目标后取得", "笼控器": "不在行商出售；『沉寂之屋』；用于黑诞/呼吸的红苔/愈创之心",
}


SOURCES = [
    {"title": "PRTS：沉沦者的黑流树海", "url": "https://prts.wiki/w/沉沦者的黑流树海", "role": "节点机制、战斗收益、特殊规则", "accessed": ACCESSED},
    {"title": "PRTS：事件一览", "url": "https://prts.wiki/w/沉沦者的黑流树海/事件一览", "role": "事件标题、选项与结果", "accessed": ACCESSED},
    {"title": "PRTS：引擎配件目录", "url": "https://prts.wiki/w/沉沦者的黑流树海/引擎配件目录", "role": "零件池、来源与交换规则", "accessed": ACCESSED},
    {"title": "PRTS：拟造物质编目", "url": "https://prts.wiki/w/沉沦者的黑流树海/拟造物质编目", "role": "收藏品目录与效果复核", "accessed": ACCESSED},
    {"title": "ArknightsGameData", "url": "https://github.com/Kengxxiao/ArknightsGameData", "role": "公开客户端数据：节点、场景、选项、关卡、零件、收藏品", "accessed": ACCESSED},
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_table_width(table, width=9360):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width)); tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        set_cell_shading(cell, "E8EEF5")
        for r in cell.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(font_size)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for p in table.rows[0].cells[0].paragraphs:
        p.paragraph_format.keep_with_next = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = clean(value)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    r.font.size = Pt(font_size)
    set_table_width(table)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def option_result_text(o):
    result = clean(o.get("result"))
    transition = clean(o.get("transition"))
    if result and transition:
        return f"{result}；{transition}"
    return result or transition or "无额外效果"


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)
    paragraph.add_run(" 页")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def configure_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Title", 26, "1F4D78", 0, 12), ("Subtitle", 12, "5B6573", 0, 8),
        ("Heading 1", 16, "2E74B5", 18, 10), ("Heading 2", 13, "1F4D78", 14, 7), ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        st = styles[name]
        st.font.name = "Calibri"; st._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "沉沦者的黑流树海 · 模拟器规则规范"
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in header.runs:
            r.font.size = Pt(8); r.font.color.rgb = RGBColor(91, 101, 115)
        add_page_number(section.footer.paragraphs[0])


def resolve_by_name(items, names):
    by_name = {x.get("name"): x.get("id") for x in items.values() if x.get("name")}
    return {name: by_name.get(name) for name in names}


def main():
    with SOURCE.open("r", encoding="utf-8") as f:
        raw = json.load(f)
    d = raw["details"]["rogue_6"]
    scrap = raw["modules"]["rogue_6"]["scrap"]
    items = d["items"]

    node_catalog = []
    for nid, n in d["nodeTypeData"].items():
        mechanics_confidence = "B" if nid in NODE_RULES else "C"
        node_catalog.append({
            "id": nid,
            "name": n.get("name"),
            "description": clean(n.get("description")),
            "simulator_rule": NODE_RULES.get(nid, "客户端存在该节点；详细触发规则尚待补充。"),
            # The client table proves the display metadata, while the semantic
            # mechanics below are documented by PRTS rather than client code.
            "display_confidence": "A",
            "mechanics_confidence": mechanics_confidence,
            "confidence": mechanics_confidence,
        })

    parts = []
    data_by_type = {"MOVE": scrap["moveScrapData"], "GOODS": scrap["goodsScrapData"], "PASSIVE": scrap["passiveScrapData"]}
    for pid, ptype in scrap["scrapItemToType"].items():
        item = items.get(pid, {})
        mech = deepcopy(data_by_type[ptype].get(pid, {}))
        parts.append({
            "id": pid, "name": item.get("name"), "type": ptype,
            "type_name": scrap["scrapTypeData"][ptype]["typeName"], "rarity": item.get("rarity"),
            "usage": clean(item.get("usage")), "description": clean(mech.get("scrapDesc")),
            "sell_price": mech.get("sellPrice"), "mechanics": mech,
            "acquisition": PART_ACQUISITION.get(item.get("name"), "通用来源：行商、作战掉落或不期而遇；具体权重未公开"),
        })

    relics = []
    for iid, item in items.items():
        if item.get("type") == "RELIC":
            relics.append({
                "id": iid, "name": item.get("name"), "rarity": item.get("rarity"),
                "usage": clean(item.get("usage")), "obtain_approach": clean(item.get("obtainApproach")),
                "can_sacrifice": item.get("canSacrifice"), "unlock_condition": clean(item.get("unlockCondDesc")),
                "buffs": d.get("relics", {}).get(iid, {}).get("buffs", []),
            })

    stages = [{k: v.get(k) for k in ("id", "code", "name", "levelId", "isBoss", "isElite", "difficulty", "specialNodeId")} for v in d["stages"].values()]
    special_names = ["他缚", "剑锤", "一串钱伥", "翱翼", "虬蜕", "猎印", "追忆", "犬植浆", "源私钥", "怦然", "焚毁文明", "α", "击坠神明", "厄运火杆", "同行者", "湮没光明", "四叶草化石"]

    payload = {
        "meta": {
            "title": "沉沦者的黑流树海：节点、事件、收益与池规则库",
            "version": "2026-09-01-audited1", "locale": "zh-CN", "topic_id": "rogue_6",
            "purpose": "供模拟器在进入节点、显示选项、结算成本与收益时使用",
            "status": "证据目录；不是可直接执行的完整服务端事件图",
            "evidence_levels": {"A": "公开客户端数据直接确认", "B": "PRTS规则/事件页确认", "C": "社区路线或页面不完整，需实测复核"},
            "critical_warning": "不得把未知概率默认成等概率；不得把未知效果当作零收益或通用奖励；不得把全收藏品目录直接当作任一随机收藏品池。",
        },
        "execution_order": [
            "校验移动目标和行动力/加工品移动条件", "触发装载加工品和概念体的到达效果", "确定节点场景；仅从满足楼层、结局、旗标、保密度条件的候选中抽取", "显示满足requirement的选项", "原子扣除cost；生命扣除规则若注明最低保留1则强制钳制", "执行result；随机结果必须引用具名池和外部权重", "若进入战斗，战后按层级及关卡类型结算", "写入事件旗标、节点完成/可重复状态", "处理出口、远征归队和下一层状态迁移"
        ],
        "required_state": ["floor", "zone_id", "action_points", "target_hp", "target_hp_max", "shield", "hope", "ingots", "squad_limit", "parts_box_capacity", "parts", "loaded_move_part", "collectibles", "operators", "stored_vouchers", "confidentiality", "life_game_organs", "event_flags", "ending_flags", "merchant_investment_max", "merchant_defeated", "revealed_nodes", "resident_nodes", "rng_seed"],
        "nodes": node_catalog,
        "events": EVENTS,
        "battle_rewards": {"by_floor": BATTLE_REWARDS, "special": SPECIAL_BATTLES},
        "pools": {
            "parts": {"catalog": parts, "generic_draw_rule": "仅从来源条件允许且已解锁的零件中抽取；权重由外部配置提供", "cultivate_seed_candidates": ["摇篮生物（具体子池未公开）", "浪花", "雾滚草", "小八界", "中/高稀有零件"]},
            "collectibles": {
                "catalog_all_client_defined": relics,
                "named_event_rewards": resolve_by_name(items, special_names),
                "random_draw_candidates_status": "UNRESOLVED_SERVER_OR_RULE_CONFIG",
                "advisory_filter_only": "可用作候选初筛：已解锁、非明确结局专属、满足can_sacrifice/场景要求；此筛选不等于官方池。",
            },
            "stages": {"catalog": stages, "node_specific_weights_status": "UNRESOLVED"},
            "event_scenes_raw": list(d["choiceScenes"].values()),
            "choices_raw": list(d["choices"].values()),
        },
        "client_counts": {"nodes": len(node_catalog), "events_manual": len(EVENTS), "choice_scenes": len(d["choiceScenes"]), "choices": len(d["choices"]), "stages": len(stages), "parts": len(parts), "collectibles": len(relics)},
        "unknowns": [
            "各楼层完整事件候选池与抽取权重未在公开客户端表中完整给出。", "随机收藏品的每个场景专属候选池/权重未完整公开。",
            "行商库存、刷新和特殊槽位的精确权重未完整公开。", "部分连续事件（如洞中宝）后续轮次和概率需要实测。",
            "第VII层常规战斗收益表在所查规则页为空。", "公开资料会更新；版本升级后应按topic表差分重建。"
        ],
        "sources": SOURCES,
    }
    with OUT_JSON.open("w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    doc = Document(); configure_doc(doc)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(110)
    run = p.add_run("沉沦者的黑流树海")
    run.bold = True; run.font.size = Pt(28); run.font.color.rgb = RGBColor.from_string("1F4D78")
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("节点、事件、收益与藏品/零件池\n模拟器实现规范")
    r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string("2E74B5")
    doc.add_paragraph()
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("版本：2026-08-31 · 对应客户端主题 rogue_6\n配套机器可读文件：黑流树海模拟器规则库.json").font.size = Pt(10)
    box = doc.add_table(rows=1, cols=1); box.alignment = WD_TABLE_ALIGNMENT.CENTER; set_table_width(box, 7000)
    box.cell(0, 0).text = "目标：让AI玩家在进入任一节点时，能够判断可发生的场景、可点击选项、前置与代价、确定收益、随机池引用及后续状态迁移。\n\n重要：未知概率不得默认等概率；“全收藏品目录”不等于任一随机收藏品池。"
    set_cell_shading(box.cell(0, 0), "E8EEF5")
    for rr in box.cell(0,0).paragraphs[0].runs: rr.font.size = Pt(11)
    doc.add_page_break()

    doc.add_heading("1. 使用范围与证据等级", level=1)
    doc.add_paragraph(f"本规范把公开客户端数据与PRTS规则页组合为一个可执行的状态机输入。它覆盖全部21类客户端节点、客户端目录中的338个场景与396个选项原文、105个关卡、30个零件和{len(relics)}件收藏品物品。客户端另有371条藏品/难度相关效果配置，不能与物品数混为一谈。人工核验部分列出当前可识别的全部事件及其选项/收益；原始客户端目录保留在配套JSON，供后续差分与补漏。")
    add_table(doc, ["等级", "含义", "模拟器策略"], [["A", "公开客户端表直接确认", "可直接实现字段和确定效果"], ["B", "PRTS规则/事件页确认", "可实现；保留来源与版本"], ["C", "社区路线、页面不完整或需实测", "仅作为候选/提示，不用于强制结算"]], [0.7, 2.3, 3.5])
    doc.add_heading("实现边界", level=2)
    add_bullet(doc, "本版能确定“有哪些节点/事件、显示哪些已知选项、成本与确定性结果、如何引用零件和收藏品目录”。")
    add_bullet(doc, "本版不能证明所有随机场景、行商库存、收藏品抽取与连续事件分支的真实权重；这些必须通过外部权重配置、抓包或大样本实测补齐。")
    add_bullet(doc, "lubaio.wiki 当前无法解析；本版未把它当作证据源，改用PRTS和公开客户端数据。")

    doc.add_heading("2. 模拟器结算顺序", level=1)
    for i, step in enumerate(payload["execution_order"], 1):
        doc.add_paragraph(f"{i}. {step}")
    doc.add_heading("最小状态字段", level=2)
    add_table(doc, ["字段组", "字段"], [
        ["地图", "floor, zone_id, action_points, revealed_nodes, resident_nodes"],
        ["资源", "target_hp, target_hp_max, shield, hope, ingots, squad_limit"],
        ["物品", "parts, loaded_move_part, parts_box_capacity, collectibles, stored_vouchers"],
        ["队伍", "operators（含晋升、临时招募、远征状态）"],
        ["规则旗标", "confidentiality, life_game_organs, event_flags, ending_flags"],
        ["商店/随机", "merchant_investment_max, merchant_defeated, rng_seed"],
    ], [1.2, 5.3])
    doc.add_paragraph("建议所有 option 采用统一结构：requirement → atomic cost → result[] → transition。随机结果必须带 pool_id；若 pool_id 的权重状态为 UNRESOLVED，模拟器应暂停抽取或使用显式标记的测试权重。")

    doc.add_heading("3. 节点总表（21类）", level=1)
    add_table(doc, ["客户端ID", "名称", "进入/结算规则"], [[n["id"], n["name"], n["simulator_rule"]] for n in node_catalog], [1.2, 1.0, 4.3], 8)

    doc.add_heading("4. 战斗基础收益", level=1)
    doc.add_paragraph("表中经验和源石锭为基础结算范围；收藏品、零件、招募券、关卡特殊掉落及藏品加成应在战后追加处理。第VII层公开规则表为空，不得沿用第VI层猜值。")
    add_table(doc, ["层", "类型", "经验", "源石锭"], BATTLE_REWARDS, [0.7, 1.4, 1.2, 1.2])
    doc.add_heading("特殊战斗收益", level=2)
    add_table(doc, ["事件/关卡", "经验", "源石锭"], SPECIAL_BATTLES, [3.4, 1.3, 1.3])

    doc.add_heading("5. 事件与选项规则", level=1)
    doc.add_paragraph("以下按事件场景逐项描述。『楼层池未公开』表示事件本身和选项已确认，但无法从现有公开数据证明其所有楼层与真实抽取权重。JSON同时保留客户端338个场景和396个选项原文，便于继续核验。")
    grouped = {}
    for ev in EVENTS: grouped.setdefault(ev["node_type"], []).append(ev)
    for node_type, evs in grouped.items():
        doc.add_heading(node_type, level=2)
        for ev in evs:
            doc.add_heading(ev["name"], level=3)
            meta_p = doc.add_paragraph(f"出现条件：{ev['availability']}　｜　证据：{ev['confidence']}　｜　可重复：{'是' if ev['repeatable'] else '否/按场景结束'}")
            meta_p.paragraph_format.keep_with_next = True
            add_table(doc, ["选项", "前置", "代价", "结果与后续"], [[o["label"], o["requirement"], o["cost"], option_result_text(o)] for o in ev["options"]], [1.25, 1.55, 1.35, 2.35], 7.8)
            if ev["notes"]:
                p = doc.add_paragraph(); p.add_run("实现注：").bold = True; p.add_run(ev["notes"])

    doc.add_heading("6. 零件池（30件）", level=1)
    doc.add_paragraph("零件分为加工品（移动）、自然物（估价/培育）和概念体（到达节点触发）。模拟器应先按来源规则过滤，再应用外部权重。一般无特殊说明的零件可来自行商、作战掉落或不期而遇；下表列出客户端确定效果与PRTS来源限制。")
    for t in ("MOVE", "GOODS", "PASSIVE"):
        tparts = [p for p in parts if p["type"] == t]
        doc.add_heading(tparts[0]["type_name"], level=2)
        add_table(doc, ["名称", "稀有度", "效果", "售价", "来源/限制"], [[p["name"], p["rarity"], p["description"] or p["usage"], p["sell_price"], p["acquisition"]] for p in tparts], [1.25, 0.65, 2.55, 0.55, 1.5], 7.7)
    doc.add_heading("培育与交换池规则", level=2)
    add_bullet(doc, "种子培育候选：一个摇篮生物、浪花、雾滚草、小八界，或中/高稀有零件；具体权重未公开。")
    add_bullet(doc, "低/中稀有零件通常交换为同稀有度；高稀有零件可换高稀有零件（排除笼控器），或降档为一次性喷气背包/涂装黎博利/涂装佩洛/涂装阿戈尔；任意零件允许换到自身。")
    add_bullet(doc, "概念体在装载加工品并移动到匹配节点时结算；一次性概念体应在触发后扣除激活次数。")

    doc.add_heading("7. 收藏品池边界", level=1)
    doc.add_paragraph(f"客户端定义的收藏品物品目录共{len(relics)}件，完整ID、名称、稀有度、效果、解锁条件、能否献祭与可匹配buff字段均在配套JSON中。这里不把{len(relics)}件目录直接列为任一『随机收藏品』池，因为事件池、商店池、战后池和结局专属池并不等价。")
    named = payload["pools"]["collectibles"]["named_event_rewards"]
    add_table(doc, ["事件物品", "客户端ID", "用途"], [[k, v or "未由同名直接解析", "按第5章对应事件/结局状态机使用"] for k, v in named.items()], [1.5, 2.4, 2.6])
    doc.add_heading("推荐池接口", level=2)
    add_table(doc, ["pool_id", "候选过滤", "权重策略"], [
        ["collectible.random.generic", "已解锁；排除明确结局专属；满足场景限制", "未解析：必须外部注入"],
        ["collectible.event.named", "事件明确指定的单品或具名集合", "确定性/按事件配置"],
        ["part.random.generic", "来源允许、已解锁、容量允许", "未解析：必须外部注入"],
        ["part.seed.cultivate", "摇篮生物/浪花/雾滚草/小八界/中高稀有零件", "未解析：必须外部注入"],
        ["stage.node.random", "关卡类型、层级、结局与节点条件匹配", "未解析：必须外部注入"],
    ], [1.8, 3.0, 1.7])

    doc.add_heading("8. 已知缺口与验收标准", level=1)
    for x in payload["unknowns"]: add_bullet(doc, x)
    doc.add_heading("模拟器不得做的事", level=2)
    add_bullet(doc, "不得把未知候选池当成全目录；不得把未知权重默认成均匀分布后标作真实规则。")
    add_bullet(doc, "不得在前置条件不满足时展示或执行选项；扣除成本与发放收益必须是原子操作。")
    add_bullet(doc, "不得忽略事件旗标、结局互斥、保密度、生命游戏器官、行商击破状态和可重复访问状态。")
    doc.add_heading("建议验收用例", level=2)
    add_table(doc, ["用例", "期望"], [
        ["进入空地", "不产生事件、战斗或收益"], ["笼控器进入黑诞", "显示专属选项并可获得猎印"],
        ["4锭进入货从口出", "可购买1次；扣4锭；从合法零件池抽1件；保留继续/离开"],
        ["狭路相逢失败", "目标生命和护盾不变；无胜利奖励"], ["险路尽头剩3行动力", "加工品+1，希望+3，行动力清零，进入下一层"],
        ["持涂装佩洛进入得偿所愿", "先结算概念体：随机收藏品+1，再显示节点选项"],
        ["随机池权重缺失", "抛出UNRESOLVED_POOL或使用显式test_profile；不得静默均匀抽样"],
    ], [2.1, 4.4])

    doc.add_heading("9. 数据来源与版本", level=1)
    doc.add_paragraph("来源按用途交叉核对。PRTS页面为社区维护资料，客户端仓库为公开解包数据；两者均可能随版本更新。模拟器应记录本规范版本并在客户端数据更新时做ID级差分。")
    add_table(doc, ["来源", "用途", "访问日期", "链接"], [[s["title"], s["role"], s["accessed"], s["url"]] for s in SOURCES], [1.45, 2.1, 0.85, 2.1], 7.5)
    doc.add_paragraph(f"配套JSON内容：21类节点、人工核验事件规则、30件零件完整机制、{len(relics)}件收藏品物品目录、105个关卡目录、338个客户端场景原文、396个客户端选项原文，以及所有未解析池的显式状态。")
    doc.save(OUT_DOCX)
    print(json.dumps({"docx": str(OUT_DOCX), "json": str(OUT_JSON), "counts": payload["client_counts"]}, ensure_ascii=False))


if __name__ == "__main__":
    main()
